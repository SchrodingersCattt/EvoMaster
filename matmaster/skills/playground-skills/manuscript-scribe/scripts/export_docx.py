"""
Export an assembled Markdown manuscript to Word (.docx) format.

Requires ``python-docx`` (install via ``pip install python-docx`` or the project's
``docx`` optional dependency group).

Features:
  - Heading mapping: # -> Title, ## -> Heading 1, ### -> Heading 2
  - Inline formatting: **bold**, *italic*, `code`
  - Citations: [n](url) -> superscript "[n]" with hyperlink
  - Bullet lists (- or *) and numbered lists (1.)
  - Basic table support (Markdown pipe tables)
  - Smart scientific formatting:
    - Chemical formula subscripts (CO2 -> CO₂ with Word subscript)
    - Physics subscripts: _{eff} -> proper Word subscript
    - Physics superscripts: ^{-6} -> proper Word superscript
    - En-dash and minus sign preservation
    - Reference list auto-formatting (journal italic, year bold)
  - Configurable font, margins, page size
  - Optional .docx style template via --style_template

Usage:
  python export_docx.py --input final.md --output manuscript.docx
  python export_docx.py --input final.md --output manuscript.docx --style_template template.docx

Output: Creates a .docx file at --output path.
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ---------------------------------------------------------------------------
# Chemical formula detection
# ---------------------------------------------------------------------------

# Common element symbols (1- and 2-letter)
_ELEMENTS = {
    'H',
    'He',
    'Li',
    'Be',
    'B',
    'C',
    'N',
    'O',
    'F',
    'Ne',
    'Na',
    'Mg',
    'Al',
    'Si',
    'P',
    'S',
    'Cl',
    'Ar',
    'K',
    'Ca',
    'Sc',
    'Ti',
    'V',
    'Cr',
    'Mn',
    'Fe',
    'Co',
    'Ni',
    'Cu',
    'Zn',
    'Ga',
    'Ge',
    'As',
    'Se',
    'Br',
    'Kr',
    'Rb',
    'Sr',
    'Y',
    'Zr',
    'Nb',
    'Mo',
    'Tc',
    'Ru',
    'Rh',
    'Pd',
    'Ag',
    'Cd',
    'In',
    'Sn',
    'Sb',
    'Te',
    'I',
    'Xe',
    'Cs',
    'Ba',
    'La',
    'Ce',
    'Pr',
    'Nd',
    'Pm',
    'Sm',
    'Eu',
    'Gd',
    'Tb',
    'Dy',
    'Ho',
    'Er',
    'Tm',
    'Yb',
    'Lu',
    'Hf',
    'Ta',
    'W',
    'Re',
    'Os',
    'Ir',
    'Pt',
    'Au',
    'Hg',
    'Tl',
    'Pb',
    'Bi',
    'Po',
    'At',
    'Rn',
    'Fr',
    'Ra',
    'Ac',
    'Th',
    'Pa',
    'U',
    'Np',
    'Pu',
    'Am',
}

# Pattern: element symbol followed by digits, repeated 2+ times
# e.g. CO2, H2O, Fe2O3, C7H8N2O, CH4
_CHEM_FORMULA_RE = re.compile(r'\b((?:[A-Z][a-z]?\d*){2,})\b')


def _is_chemical_formula(text: str) -> bool:
    """Check if text looks like a chemical formula (Element+optional_count, 2+ elements)."""
    # Parse into element-count pairs
    pairs = re.findall(r'([A-Z][a-z]?)(\d*)', text)
    if not pairs:
        return False
    # Filter out empty matches at the end
    pairs = [(elem, count) for elem, count in pairs if elem]
    if len(pairs) < 2:
        return False
    # Check that all "elements" are real element symbols
    return all(elem in _ELEMENTS for elem, _ in pairs)


def _add_chemical_formula(paragraph, formula: str) -> None:
    """Add a chemical formula with subscripted numbers to a paragraph."""
    pairs = re.findall(r'([A-Z][a-z]?)(\d*)', formula)
    for elem, count in pairs:
        if not elem:
            continue
        run = paragraph.add_run(elem)
        if count:
            run = paragraph.add_run(count)
            run.font.subscript = True


# ---------------------------------------------------------------------------
# CJK spacing fix
# ---------------------------------------------------------------------------

# CJK unified ideographs + extensions + compatibility
_CJK_CHAR = (
    r'[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff'
    r'\u3000-\u303f\uff00-\uffef'
    r'\u2e80-\u2eff\u3100-\u312f]'
)

_CJK_SPACE_AFTER = re.compile(
    rf"({_CJK_CHAR})\s+([\d(A-Za-z\u2212\u2013\u2014\u00b1\u2248\u2264\u2265<>≈±])"
)
_CJK_SPACE_BEFORE = re.compile(
    rf"([\d)A-Za-z%°\u2212\u2013\u00b1\u2248\u2264\u2265])\s+({_CJK_CHAR})"
)


def _fix_cjk_spacing(text: str) -> str:
    """Remove spurious spaces between CJK characters and adjacent numbers/units.

    Rules:
    * CJK + space + digit/latin → remove space  (约 1.9 → 约1.9)
    * digit/latin + space + CJK → remove space  (eV 的  → eV的)
    * Latin + space + Latin is preserved          (1.9 eV stays)
    """
    text = _CJK_SPACE_AFTER.sub(r'\1\2', text)
    text = _CJK_SPACE_BEFORE.sub(r'\1\2', text)
    return text


# ---------------------------------------------------------------------------
# Markdown escape stripping (for science notation)
# ---------------------------------------------------------------------------


def _strip_md_escapes(text: str) -> str:
    r"""Remove Markdown backslash escapes that break science notation.

    ``\_\{g\}``  →  ``_{g}``
    ``\*E\*``    →  ``*E*``
    Only strips escapes before ``_``, ``{``, ``}``, ``*``, ``^``, ``-``.
    """
    return re.sub(r'\\([_{}\-*^])', r'\1', text)


# ---------------------------------------------------------------------------
# Smart inline formatting: subscripts, superscripts, chemical formulas
# ---------------------------------------------------------------------------


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add a superscript hyperlink run to a paragraph (for citations)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)

    run_elem.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)


def _add_plain_text_with_science(paragraph, text: str) -> None:
    """Add plain text with auto-detection of chemical formulas and
    subscript/superscript notation.

    Handles:
      - _{text} -> Word subscript  (also after Markdown escape stripping)
      - ^{text} -> Word superscript
      - Chemical formulas (CO2, H2O, Fe2O3) -> element + subscript numbers
      - CJK spacing fix (remove spurious spaces between CJK and digits/units)
    """
    # Strip Markdown backslash escapes that hide science notation
    text = _strip_md_escapes(text)
    # Fix CJK spacing before emitting runs
    text = _fix_cjk_spacing(text)

    # Pattern for sub/superscript notation and potential chemical formulas
    pattern = re.compile(
        r'(_\{([^}]+)\})'  # _{text} subscript
        r'|(\^\{([^}]+)\})'  # ^{text} superscript
        r'|(\b[A-Z][a-z]?(?:\d+)[A-Z]?[a-z]?(?:\d+)?(?:[A-Z][a-z]?(?:\d+)?)*\b)'  # potential chemical formula
    )

    pos = 0
    for m in pattern.finditer(text):
        # Plain text before match
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])

        if m.group(2) is not None:  # subscript _{text}
            run = paragraph.add_run(m.group(2))
            run.font.subscript = True
        elif m.group(4) is not None:  # superscript ^{text}
            run = paragraph.add_run(m.group(4))
            run.font.superscript = True
        elif m.group(5) is not None:  # potential chemical formula
            candidate = m.group(5)
            if _is_chemical_formula(candidate):
                _add_chemical_formula(paragraph, candidate)
            else:
                paragraph.add_run(candidate)

        pos = m.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_formatted_text(paragraph, text: str, is_reference_entry: bool = False) -> None:
    """Parse inline Markdown and add runs with formatting to a paragraph.

    Handles: **bold**, *italic*, `code`, [n](url) citations, _{sub}, ^{sup},
    chemical formulas (CO2, H2O, etc.), CJK spacing fixes.

    If is_reference_entry=True, applies special formatting: journal names italic,
    year bold, page-range en-dash.
    """
    # Strip Markdown backslash escapes early so \_\{g\} becomes _{g}
    text = _strip_md_escapes(text)

    if is_reference_entry:
        _add_reference_entry(paragraph, text)
        return

    # Master pattern: process in priority order
    # Citations first (to avoid [n] being caught by other patterns),
    # then bold, italic, code, then everything else via plain text handler
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'  # bold
        r'|(\*([^*]+?)\*)'  # italic (single *)
        r'|(`([^`]+?)`)'  # code
        r'|(\[(\d+)\]\(([^)]+)\))'  # citation [n](url)
        r'|(\[(\d+)\])'  # plain citation [n]
    )

    pos = 0
    for m in pattern.finditer(text):
        # Add plain text before this match (with science formatting)
        if m.start() > pos:
            _add_plain_text_with_science(paragraph, text[pos : m.start()])

        if m.group(2):  # bold  (groups 1,2)
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(4):  # italic  (groups 3,4)
            run = paragraph.add_run(m.group(4))
            run.italic = True
        elif m.group(6):  # code  (groups 5,6)
            run = paragraph.add_run(m.group(6))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif m.group(8):  # citation with url [n](url)  (groups 7,8,9)
            _add_hyperlink(paragraph, f"[{m.group(8)}]", m.group(9))
        elif m.group(10):  # plain citation [n]  (groups 10,11)
            run = paragraph.add_run(f"[{m.group(11)}]")
            run.font.superscript = True

        pos = m.end()

    # Add remaining plain text (with science formatting)
    if pos < len(text):
        _add_plain_text_with_science(paragraph, text[pos:])


# ---------------------------------------------------------------------------
# Reference entry formatting
# ---------------------------------------------------------------------------


def _add_reference_entry(paragraph, text: str) -> None:
    """Format a reference entry: journal italic, year bold, en-dash for pages.

    Expected input: "[n] Authors. Title. *Journal*, **Year**, Volume, Pages. URL"
    or:             "[n] *Journal*, **Year**, Volume, Pages."
    """
    # Handle the [n] prefix
    ref_match = re.match(r'^\[(\d+)\]\s*', text)
    if ref_match:
        run = paragraph.add_run(f"[{ref_match.group(1)}] ")
        text = text[ref_match.end() :]

    # Now process the rest with standard formatting
    # Pattern for bold (year), italic (journal), and URLs
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'  # bold (year)
        r'|(\*([^*]+?)\*)'  # italic (journal)
        r'|(https?://[^\s]+)'  # URL
    )

    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            # Plain text -- replace hyphen in page ranges with en-dash
            plain = text[pos : m.start()]
            plain = re.sub(r'(\d+)-(\d+)', r'\1–\2', plain)
            paragraph.add_run(plain)

        if m.group(2):  # bold (year)
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(4):  # italic (journal)
            run = paragraph.add_run(m.group(4))
            run.italic = True
        elif m.group(5):  # URL
            url = m.group(5).rstrip('.,;:)')
            run = paragraph.add_run(url)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(5, 99, 193)

        pos = m.end()

    if pos < len(text):
        plain = text[pos:]
        plain = re.sub(r'(\d+)-(\d+)', r'\1–\2', plain)
        paragraph.add_run(plain)


# ---------------------------------------------------------------------------
# Block-level parsing
# ---------------------------------------------------------------------------


def _is_table_line(line: str) -> bool:
    stripped = line.strip()
    return (
        stripped.startswith('|') and stripped.endswith('|') and stripped.count('|') >= 2
    )


def _is_separator_line(line: str) -> bool:
    return bool(re.match(r'^\s*\|[\s\-:|]+\|\s*$', line))


def _parse_table_row(line: str) -> list[str]:
    cells = line.strip().strip('|').split('|')
    return [c.strip() for c in cells]


def _add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = ''
                _add_formatted_text(cell.paragraphs[0], cell_text)
    if rows:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True


def _is_in_references_section(lines: list[str], current_idx: int) -> bool:
    """Check if current line is inside a ## References section."""
    for j in range(current_idx, -1, -1):
        stripped = lines[j].strip()
        if re.match(r'^##\s+', stripped):
            return 'reference' in stripped.lower()
    return False


def export_markdown_to_docx(
    md_text: str,
    output_path: str | Path,
    style_template: str | Path | None = None,
) -> None:
    """Convert Markdown text to a .docx file."""
    if not HAS_DOCX:
        print(
            'Error: python-docx is not installed. '
            'Install it with: pip install python-docx',
            file=sys.stderr,
        )
        sys.exit(1)

    if style_template and Path(style_template).exists():
        doc = Document(str(style_template))
    else:
        doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    lines = md_text.splitlines()
    i = 0
    in_references = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip HTML comments
        if stripped.startswith('<!--') and stripped.endswith('-->'):
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('### '):
            p = doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith('## '):
            sec_name = stripped[3:].strip()
            in_references = 'reference' in sec_name.lower()
            p = doc.add_heading(sec_name, level=1)
            i += 1
            continue
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue

        # Table
        if _is_table_line(stripped):
            table_rows: list[list[str]] = []
            while i < len(lines) and _is_table_line(lines[i].strip()):
                if not _is_separator_line(lines[i]):
                    table_rows.append(_parse_table_row(lines[i]))
                i += 1
            _add_table(doc, table_rows)
            continue

        # Bullet list
        if re.match(r'^\s*[-*]\s+', stripped):
            text = re.sub(r'^\s*[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\s*\d+\.\s+', stripped):
            text = re.sub(r'^\s*\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)
            i += 1
            continue

        # Figure placeholder
        if stripped.lower().startswith('figure ') and '.' in stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Reference entry: line starting with [n]
        if in_references and re.match(r'^\[\d+\]', stripped):
            p = doc.add_paragraph()
            # Collect multi-line reference entry
            ref_lines = [stripped]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if (
                    not next_line
                    or re.match(r'^\[\d+\]', next_line)
                    or next_line.startswith('#')
                ):
                    break
                ref_lines.append(next_line)
                i += 1
            full_ref = ' '.join(ref_lines)
            _add_formatted_text(p, full_ref, is_reference_entry=True)
            continue

        # Regular paragraph
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line:
                break
            if next_line.startswith('#') or _is_table_line(next_line):
                break
            if re.match(r'^\s*[-*]\s+', next_line) or re.match(
                r'^\s*\d+\.\s+', next_line
            ):
                break
            if next_line.startswith('<!--'):
                break
            if in_references and re.match(r'^\[\d+\]', next_line):
                break
            para_lines.append(next_line)
            i += 1

        full_para = ' '.join(para_lines)
        p = doc.add_paragraph()
        _add_formatted_text(p, full_para)

    # Page numbers in footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._element.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Word document exported to {output_path}.")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main() -> None:
    ap = argparse.ArgumentParser(
        description='Export assembled Markdown manuscript to Word (.docx) format.'
    )
    ap.add_argument('--input', required=True, help='Path to assembled Markdown file')
    ap.add_argument('--output', required=True, help='Path to output .docx file')
    ap.add_argument(
        '--style_template',
        default=None,
        help='Optional .docx template file for styles (fonts, headers, etc.)',
    )
    args = ap.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    md_text = input_path.read_text(encoding='utf-8')
    export_markdown_to_docx(md_text, args.output, args.style_template)


if __name__ == '__main__':
    main()
